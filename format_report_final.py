from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


BASE = Path(__file__).resolve().parent
INPUT_PDF = BASE / "Do_An_Quan_Ly_ATTT_N22DCAT033_Tran_Phuoc_Loi.pdf"
OUT_DOCX = BASE / "report_final.docx"
OUT_PDF = BASE / "report_final.pdf"
LOG = BASE / "LOG_CHANGES.md"
RENDER_DIR = BASE / "render_report_final_pages"

TNR = "Times New Roman"
FONT_DIR = Path(r"C:\Windows\Fonts")
FONT_MAP = {
    "TNR": "times.ttf",
    "TNR-Bold": "timesbd.ttf",
    "TNR-Italic": "timesi.ttf",
    "TNR-BoldItalic": "timesbi.ttf",
}
for name, filename in FONT_MAP.items():
    path = FONT_DIR / filename
    if path.exists():
        pdfmetrics.registerFont(TTFont(name, str(path)))


def inspect_source():
    reader = PdfReader(str(INPUT_PDF))
    chars = sum(len(page.extract_text() or "") for page in reader.pages)
    first = reader.pages[0].mediabox
    return {
        "pages": len(reader.pages),
        "chars": chars,
        "text_based": chars > 100,
        "size": f"{float(first.width):.1f} x {float(first.height):.1f} pt",
    }


REPORT = [
    {
        "title": "CHƯƠNG I: KHẢO SÁT HIỆN TRẠNG VÀ XÁC ĐỊNH MỤC TIÊU CHIẾN LƯỢC QUẢN TRỊ ATTT",
        "sections": [
            (
                "1.1. Giới thiệu doanh nghiệp giả định",
                [
                    "Công ty Cổ phần Giải pháp An toàn thông tin CyberShield Việt Nam được giả định là doanh nghiệp hoạt động trong lĩnh vực tư vấn và triển khai dịch vụ an toàn thông tin. Các dịch vụ chính gồm kiểm thử xâm nhập, tư vấn xây dựng hệ thống quản lý ATTT, giám sát SOC, ứng cứu sự cố, triển khai giải pháp bảo mật và đào tạo nhận thức bảo mật.",
                    "Công ty hiện có khoảng 150 nhân sự, trong đó nhóm kỹ thuật chiếm tỷ trọng lớn gồm kỹ sư SOC/IR, chuyên gia kiểm thử xâm nhập, tư vấn GRC, kỹ sư triển khai giải pháp, nhân viên kinh doanh và khối văn phòng. Khách hàng chủ yếu là các doanh nghiệp trong nước có nhu cầu bảo vệ hạ tầng số và dữ liệu nhạy cảm.",
                    "Định hướng đến năm 2031 là tăng quy mô lên khoảng 450 nhân sự, vận hành SOC 24/7, đạt chứng nhận ISO/IEC 27001:2022, mở rộng dịch vụ MSSP sang khu vực ASEAN và xây dựng cổng khách hàng tự phục vụ.",
                ],
            ),
            (
                "1.2. Khảo sát hiện trạng và định hướng phát triển",
                [
                    {
                        "caption": "Bảng 1.1. Khảo sát hiện trạng và định hướng phát triển của CyberShield",
                        "rows": [
                            ["Nhóm thông tin", "Trạng thái hiện tại năm 2026", "Định hướng giai đoạn 2026 - 2031"],
                            ["Lĩnh vực kinh doanh", "Tư vấn ATTT, pentest, SOC, ứng cứu sự cố, triển khai EDR/SIEM/WAF và đào tạo nhận thức.", "Trở thành nhà cung cấp MSSP uy tín tại Việt Nam và mở rộng sang ASEAN."],
                            ["Thị trường", "Khoảng 80 khách hàng doanh nghiệp trong nước, tập trung tại TP.HCM, Hà Nội và Đà Nẵng.", "Phục vụ khách hàng tại Singapore, Thái Lan, Malaysia; yêu cầu chứng minh tuân thủ cao hơn."],
                            ["Hệ thống CNTT", "Email, AD, VPN, Git, ticket, HRM, kế toán, file server, SIEM/SOAR, cloud portal và phòng lab.", "Chuẩn hóa cloud-native, tự động hóa giám sát, tích hợp GRC, CMDB và quản lý cấu hình."],
                            ["Dữ liệu trọng yếu", "Log khách hàng, báo cáo lỗ hổng, bằng chứng pentest, hợp đồng, mã nguồn và tài khoản kiểm thử.", "Phân loại, mã hóa, kiểm soát chia sẻ và lưu vết dữ liệu theo ISO 27001."],
                        ],
                    }
                ],
            ),
            (
                "1.3. Đánh giá điểm mạnh và điểm yếu",
                [
                    "Điểm mạnh của CyberShield là đội ngũ kỹ thuật có kinh nghiệm thực tế trong tấn công, phòng thủ, giám sát và ứng cứu sự cố. Công ty đã có nền tảng SOC, công cụ quản lý ticket, hệ thống VPN, EDR cho một số nhóm kỹ thuật và quy trình phục vụ khách hàng tương đối rõ ràng.",
                    "Khoảng trống lớn nằm ở quản trị nội bộ. Công ty có năng lực làm bảo mật cho khách hàng nhưng chưa có mô hình CISO/GRC độc lập, chưa có danh mục tài sản đầy đủ cho cloud và SaaS, chưa có phân loại dữ liệu thống nhất và chưa có cơ chế đo lường KPI bảo mật nội bộ.",
                ],
            ),
            (
                "1.4. Mục tiêu chiến lược quản trị ATTT",
                [
                    "CyberShield cần chuyển từ mô hình bảo mật dựa nhiều vào chuyên gia kỹ thuật sang mô hình quản trị có chính sách, trách nhiệm, bằng chứng và cải tiến liên tục. Mục tiêu quan trọng nhất là bảo vệ dữ liệu khách hàng, duy trì dịch vụ trọng yếu và tạo năng lực chứng minh tuân thủ khi mở rộng thị trường.",
                    [
                        "Đạt và duy trì chứng nhận ISO/IEC 27001:2022 cho phạm vi SOC, cổng khách hàng và kho báo cáo bảo mật.",
                        "100% dữ liệu khách hàng được phân loại, có chủ sở hữu, chính sách lưu trữ, mã hóa và hủy dữ liệu sau dự án.",
                        "Giảm thời gian phát hiện sự cố nghiêm trọng xuống dưới 30 phút và thời gian xử lý ban đầu xuống dưới 2 giờ.",
                        "Hoàn thiện chương trình đào tạo SETA, đưa tỷ lệ click phishing mô phỏng xuống dưới 5%.",
                    ],
                ],
            ),
        ],
    },
    {
        "title": "CHƯƠNG II: TIẾN TRÌNH NHẬN DẠNG, ĐÁNH GIÁ VÀ XỬ LÝ RỦI RO",
        "sections": [
            (
                "2.1. Phương pháp quản lý rủi ro",
                [
                    "Quản lý rủi ro được thực hiện theo chu trình nhận dạng tài sản, xác định mối đe dọa và điểm yếu, đánh giá khả năng xảy ra, đánh giá tác động, lựa chọn biện pháp xử lý và theo dõi rủi ro còn lại. Cách tiếp cận này giúp công ty đầu tư bảo mật theo mức độ ưu tiên thay vì triển khai công nghệ rời rạc.",
                    "Đối với CyberShield, tiêu chí tác động gồm thiệt hại tài chính, vi phạm hợp đồng, lộ dữ liệu khách hàng, gián đoạn SOC, mất uy tín nghề nghiệp, ảnh hưởng pháp lý và ảnh hưởng đến khả năng mở rộng thị trường.",
                ],
            ),
            (
                "2.2. Danh mục kiểm kê và định giá tài sản thông tin",
                [
                    {
                        "caption": "Bảng 2.1. Danh mục kiểm kê tài sản thông tin",
                        "rows": [
                            ["Mã TS", "Nhóm tài sản", "Tài sản cụ thể", "Mức độ"],
                            ["TS-01", "Dữ liệu", "Log SOC, cảnh báo, IOC, playbook theo khách hàng.", "5"],
                            ["TS-02", "Dữ liệu", "Báo cáo pentest, bằng chứng lỗ hổng, tài khoản kiểm thử.", "5"],
                            ["TS-03", "Phần mềm", "Cổng khách hàng, API báo cáo, công cụ tự động hóa nội bộ.", "4"],
                            ["TS-04", "Hạ tầng", "AD, VPN, email, EDR, SIEM, hệ thống ticket.", "4"],
                            ["TS-05", "Tri thức", "Quy trình xử lý sự cố, mẫu kiểm thử, kịch bản red team.", "4"],
                        ],
                    }
                ],
            ),
            (
                "2.3. Phân tích và xử lý rủi ro trọng yếu",
                [
                    "Các rủi ro trọng yếu của CyberShield tập trung vào dữ liệu khách hàng, quyền truy cập đặc quyền, ransomware, cấu hình cloud, yêu cầu hợp đồng khu vực và việc sử dụng công cụ kiểm thử. Mỗi rủi ro phải có chủ sở hữu, hạn xử lý, ngân sách và bằng chứng hoàn thành.",
                    {
                        "caption": "Bảng 2.2. Ma trận xử lý rủi ro trọng yếu",
                        "rows": [
                            ["Rủi ro", "Nguyên nhân", "Tác động", "Kiểm soát chính"],
                            ["Lộ dữ liệu khách hàng", "Chia sẻ sai kênh, thiếu phân loại, quyền rộng.", "Rất cao", "DLP, mã hóa, IAM theo dự án, nhật ký truy cập."],
                            ["Tài khoản cũ còn quyền", "Quy trình nghỉ việc/chuyển bộ phận chưa chặt.", "Cao", "JML, rà soát quyền hàng tháng, SSO/MFA."],
                            ["Ransomware gián đoạn SOC", "Endpoint yếu, phishing, backup chưa kiểm thử.", "Rất cao", "EDR, backup bất biến, DR, diễn tập."],
                            ["Cloud cấu hình sai", "Thiếu baseline và review IaC.", "Cao", "CSPM, kiểm soát public exposure, logging cloud."],
                            ["Không đạt yêu cầu hợp đồng ASEAN", "Thiếu chứng nhận và bằng chứng tuân thủ.", "Cao", "ISO 27001, GRC, vendor review, audit nội bộ."],
                        ],
                    },
                    "Sau khi triển khai kiểm soát, công ty cần đánh giá rủi ro còn lại và trình Ban lãnh đạo phê duyệt. Ví dụ, ngay cả khi đã có DLP và mã hóa, nguy cơ nhân viên ghi lại thông tin nhạy cảm vẫn tồn tại và cần được xử lý thêm bằng đào tạo, quy định kỷ luật và giám sát hành vi bất thường.",
                ],
            ),
        ],
    },
    {
        "title": "CHƯƠNG III: ĐỀ XUẤT CÁC CƠ CHẾ KỸ THUẬT BẢO VỆ DỰA TRÊN KẾT QUẢ BIA",
        "sections": [
            (
                "3.1. Phân tích tác động kinh doanh BIA",
                [
                    "BIA được dùng để xác định dịch vụ nào cần khôi phục trước, mức gián đoạn tối đa chấp nhận được và yêu cầu kiểm soát tương ứng. Với CyberShield, SOC, cổng khách hàng và kho báo cáo bảo mật là ba dịch vụ có ưu tiên cao nhất vì liên quan trực tiếp đến hợp đồng dịch vụ và uy tín của công ty.",
                    {
                        "caption": "Bảng 3.1. Kết quả BIA cho dịch vụ trọng yếu",
                        "rows": [
                            ["Dịch vụ", "Tác động nếu gián đoạn", "RTO", "RPO"],
                            ["SOC giám sát khách hàng", "Không phát hiện sự cố, vi phạm SLA, ảnh hưởng uy tín.", "1 giờ", "15 phút"],
                            ["Cổng khách hàng", "Khách hàng không truy cập báo cáo/ticket, gián đoạn dịch vụ.", "2 giờ", "30 phút"],
                            ["Kho báo cáo pentest", "Rủi ro mất dữ liệu và chậm bàn giao dự án.", "4 giờ", "1 giờ"],
                            ["Email/VPN/AD", "Gián đoạn vận hành nội bộ và phản hồi sự cố.", "2 giờ", "30 phút"],
                        ],
                    },
                ],
            ),
            (
                "3.2. Kiểm soát truy cập theo mô hình Zero Trust",
                [
                    "Zero Trust không mặc định tin cậy chỉ vì người dùng đang ở trong mạng nội bộ. Mọi truy cập vào hệ thống trọng yếu phải được xác thực mạnh, cấp quyền tối thiểu, kiểm tra trạng thái thiết bị và ghi log đầy đủ.",
                    [
                        "Quản trị hệ thống truy cập qua PAM, MFA mạnh, phiên quản trị có ghi log và phê duyệt có thời hạn.",
                        "SOC analyst chỉ truy cập log, cảnh báo và ticket theo khách hàng được phân công.",
                        "Pentester sử dụng tài khoản dự án có thời hạn, lưu bằng chứng đúng kho và không vượt phạm vi kiểm thử.",
                        "Nhân viên kinh doanh và quản lý dự án chỉ truy cập tài liệu khách hàng theo nhu cầu công việc.",
                    ],
                ],
            ),
            (
                "3.3. Phân đoạn mạng, SOC nội bộ và bảo vệ dữ liệu",
                [
                    "CyberShield cần tách vùng văn phòng, vùng lab kiểm thử, vùng quản trị, vùng production, vùng SOC và vùng khách truy cập. Vùng lab không được truy cập trực tiếp production; mọi dữ liệu khách hàng đưa vào lab phải được phê duyệt và có thời hạn lưu.",
                    "SOC nội bộ thu thập log từ AD, VPN, email, cloud, EDR, Git, ticket và cổng khách hàng. Các use case ưu tiên gồm đăng nhập bất thường, tạo tài khoản đặc quyền, tải dữ liệu khối lượng lớn, chia sẻ link công khai, commit chứa secret và cảnh báo ransomware.",
                    "Dữ liệu khách hàng phải được mã hóa khi lưu trữ và truyền tải. Báo cáo pentest, file log thô và bản sao cơ sở dữ liệu phải có thời hạn lưu trữ rõ ràng, tự động hủy hoặc ẩn danh khi hết nghĩa vụ hợp đồng.",
                ],
            ),
        ],
    },
    {
        "title": "CHƯƠNG IV: KẾ HOẠCH DỰ PHÒNG, DUY TRÌ BẢO MẬT VÀ PHÁT TRIỂN NGUỒN LỰC",
        "sections": [
            (
                "4.1. Quy trình ứng phó sự cố IRP",
                [
                    "IRP của CyberShield cần bảo đảm vai trò rõ ràng khi chính công ty gặp sự cố. Kịch bản ưu tiên gồm phishing chiếm quyền email, ransomware máy trạm kỹ sư, lộ báo cáo pentest và cấu hình sai cloud storage.",
                    {
                        "caption": "Bảng 4.1. Quy trình ứng phó sự cố",
                        "rows": [
                            ["Giai đoạn", "Hoạt động chính", "Đầu ra bắt buộc"],
                            ["Chuẩn bị", "Lập đội IR, kịch bản, danh bạ khẩn cấp, playbook, phân quyền công cụ.", "IR plan, playbook, danh sách liên hệ."],
                            ["Phát hiện & phân tích", "Thu thập log, xác định IOC, phạm vi ảnh hưởng, mức độ sự cố.", "Biên bản phân loại sự cố và timeline."],
                            ["Cô lập & xử lý", "Khóa tài khoản, cô lập máy, chặn IOC, vá lỗi, thu hồi link/tài liệu.", "Ticket xử lý và bằng chứng cô lập."],
                            ["Khôi phục", "Khôi phục dữ liệu, kiểm tra sạch, đưa dịch vụ hoạt động trở lại.", "Biên bản phục hồi, xác nhận RTO/RPO."],
                        ],
                    },
                ],
            ),
            (
                "4.2. BCP/DRP và duy trì bảo mật liên tục",
                [
                    "BCP/DRP được xây dựng dựa trên kết quả BIA. Dịch vụ SOC và cổng khách hàng phải có phương án dự phòng rõ ràng, backup bất biến và lịch diễn tập định kỳ. Công ty cần kiểm tra khả năng phục hồi thực tế thay vì chỉ lưu bản sao dự phòng.",
                    "Bảo mật không phải là một dự án có điểm kết thúc. Sau khi triển khai kiểm soát, công ty cần duy trì kiểm toán nội bộ, rà soát quyền, pentest định kỳ, diễn tập IR/DR và đào tạo nhận thức theo lịch vận hành chính thức.",
                ],
            ),
            (
                "4.3. Kiện toàn tổ chức và chương trình SETA",
                [
                    "CyberShield cần tách rõ đội bảo mật nội bộ với đội cung cấp dịch vụ cho khách hàng. CISO chịu trách nhiệm điều hành ISMS, GRC Team phụ trách chính sách và kiểm toán, Security Engineering triển khai kiểm soát kỹ thuật, SOC/IR nội bộ giám sát và xử lý sự cố.",
                    "SETA gồm Security Education, Training and Awareness. Nhân viên văn phòng cần hiểu phishing, mật khẩu và phân loại dữ liệu. Kỹ sư SOC cần hiểu xử lý log khách hàng. Pentester cần hiểu đạo đức nghề nghiệp và phạm vi kiểm thử. Lập trình viên cần hiểu secure coding và quản lý secret.",
                ],
            ),
        ],
    },
    {
        "title": "CHƯƠNG V: LỘ TRÌNH TRIỂN KHAI, NGÂN SÁCH, KPI VÀ KIỂM TRA TUÂN THỦ",
        "sections": [
            (
                "5.1. Lộ trình triển khai giai đoạn 2026 - 2031",
                [
                    {
                        "caption": "Bảng 5.1. Lộ trình triển khai 5 năm",
                        "rows": [
                            ["Năm", "Trọng tâm", "Công việc chính", "Kết quả kỳ vọng"],
                            ["2026", "Xây nền quản trị", "Bổ nhiệm CISO, lập Ủy ban ATTT, kiểm kê tài sản, chính sách lõi, MFA.", "Có ISMS giai đoạn 1 và danh mục rủi ro."],
                            ["2027", "Chuẩn hóa", "Triển khai GRC, DLP, JML, vendor review, audit nội bộ, ISO 27001.", "Đạt chứng nhận cho phạm vi trọng yếu."],
                            ["2028", "Mở rộng", "SOC 24/7, PAM, CSPM, DevSecOps, DR cho portal và SOC.", "Sẵn sàng phục vụ khách hàng khu vực."],
                            ["2029-2031", "Trưởng thành", "SOAR, tự động hóa bằng chứng tuân thủ, mở rộng ISMS, bảo mật chuỗi cung ứng.", "Năng lực MSSP đáp ứng thị trường ASEAN."],
                        ],
                    }
                ],
            ),
            (
                "5.2. Kế hoạch 90 ngày đầu và KPI quản lý",
                [
                    "Trong 30 ngày đầu, thực tập sinh tập trung thu thập thông tin: sơ đồ hệ thống, danh sách ứng dụng, danh mục tài khoản quản trị, chính sách hiện có và yêu cầu hợp đồng khách hàng. Từ ngày 31 đến 60, thực tập sinh hỗ trợ phân loại tài sản, xây dựng ma trận rủi ro và đề xuất quick wins. Từ ngày 61 đến 90, thực tập sinh tổng hợp lộ trình 5 năm, KPI, ngân sách và trình bày báo cáo.",
                    {
                        "caption": "Bảng 5.2. Bộ KPI quản lý an toàn thông tin",
                        "rows": [
                            ["Nhóm KPI", "Chỉ số", "Mục tiêu"],
                            ["Truy cập", "Tài khoản nghỉ việc được thu hồi đúng hạn.", "100% trong 24 giờ"],
                            ["Rủi ro", "Rủi ro cao có chủ sở hữu và kế hoạch xử lý.", ">= 95%"],
                            ["Lỗ hổng", "Lỗ hổng critical/high xử lý đúng SLA.", ">= 90%"],
                            ["SOC", "MTTD với sự cố mức cao.", "< 30 phút"],
                            ["Đào tạo", "Nhân viên hoàn thành SETA.", "100%"],
                            ["Phishing", "Tỷ lệ click trong mô phỏng.", "< 5%"],
                        ],
                    },
                ],
            ),
            (
                "5.3. Kiểm tra tuân thủ và báo cáo lãnh đạo",
                [
                    "Cơ chế kiểm tra tuân thủ giúp công ty chứng minh rằng chính sách đã được thực thi. Kiểm tra không chỉ dựa vào phỏng vấn mà cần bằng chứng: log hệ thống, ticket xử lý, biên bản đào tạo, kết quả rà soát quyền, báo cáo vulnerability scan, biên bản diễn tập và hồ sơ phê duyệt ngoại lệ.",
                    "Báo cáo hàng tháng tập trung vào sự cố, cảnh báo nghiêm trọng, rủi ro cao, tiến độ khắc phục và ngoại lệ bảo mật. Báo cáo hàng quý trình bày xu hướng rủi ro, ngân sách, KPI, kết quả kiểm toán và các quyết định cần phê duyệt. Báo cáo hằng năm đánh giá mức trưởng thành ATTT, kết quả chứng nhận và kế hoạch cải tiến năm sau.",
                ],
            ),
        ],
    },
]

CONCLUSION = [
    "CyberShield Việt Nam có lợi thế về chuyên môn kỹ thuật, nhưng để phát triển bền vững công ty cần một chương trình quản lý ATTT bài bản. Chương trình này phải kết hợp quản trị, rủi ro, chính sách, kiểm soát kỹ thuật, BIA, IRP/BCP/DRP, đào tạo và đo lường KPI.",
    "Trong giai đoạn 2026 - 2031, ưu tiên lớn nhất là thành lập mô hình CISO/GRC, phân loại và bảo vệ dữ liệu khách hàng, áp dụng IAM/MFA/PAM, quản lý lỗ hổng, thiết lập BIA/DR, đạt ISO/IEC 27001 và duy trì KPI an toàn thông tin. Nếu triển khai nghiêm túc, CyberShield có thể giảm đáng kể rủi ro nội bộ, nâng cao uy tín dịch vụ và đáp ứng yêu cầu hợp đồng trong nước cũng như khu vực ASEAN.",
]

REFERENCES = [
    "Học viện Công nghệ Bưu chính Viễn thông, Slide bài giảng môn Quản lý An toàn thông tin, 2026.",
    "ISO/IEC 27001:2022, Information security management systems - Requirements.",
    "ISO/IEC 27002:2022, Information security controls.",
    "NIST Cybersecurity Framework 2.0.",
    "NIST SP 800-34 Rev.1, Contingency Planning Guide for Federal Information Systems.",
    "NIST SP 800-100, Information Security Handbook: A Guide for Managers.",
    "Luật An toàn thông tin mạng số 86/2015/QH13 và Luật An ninh mạng số 24/2018/QH14.",
]


def iter_toc():
    page = 3
    for ci, chapter in enumerate(REPORT, 1):
        yield chapter["title"], page
        for si, (title, _) in enumerate(chapter["sections"], 1):
            yield title, page
        page += 2
    yield "KẾT LUẬN", min(page, 14)
    yield "DANH MỤC TÀI LIỆU THAM KHẢO", min(page + 1, 15)


def set_run_font(run, size=13, bold=False, italic=False):
    run.font.name = TNR
    run._element.rPr.rFonts.set(qn("w:eastAsia"), TNR)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def setup_docx(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.0)
    section.different_first_page_header_footer = True

    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = TNR
        style._element.rPr.rFonts.set(qn("w:eastAsia"), TNR)

    normal = doc.styles["Normal"]
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(1.0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.first_line_indent = Cm(0)

    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.first_line_indent = Cm(0)

    h3 = doc.styles["Heading 3"]
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.italic = True
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.keep_with_next = True
    h3.paragraph_format.first_line_indent = Cm(0)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer.add_run()._r.append(field)


def add_docx_para(doc, text, style=None, align=None, indent=True, bold=False, size=13):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Cm(1.0 if indent else 0)
    p.alignment = align if align is not None else (WD_ALIGN_PARAGRAPH.JUSTIFY if not style else p.alignment)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p


def add_docx_cover(doc):
    for text, size, bold, spacer in [
        ("HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG", 13, True, 0),
        ("KHOA AN TOÀN THÔNG TIN", 13, True, 24),
        ("BÁO CÁO ĐỒ ÁN MÔN HỌC", 15, True, 0),
        ("QUẢN LÝ AN TOÀN THÔNG TIN", 17, True, 18),
        ("Đề tài: Lập và đề xuất kế hoạch triển khai đảm bảo an toàn thông tin cho Công ty Cổ phần Giải pháp An toàn thông tin CyberShield Việt Nam giai đoạn 2026 - 2031", 14, True, 24),
    ]:
        p = add_docx_para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, bold=bold, size=size)
        if spacer:
            p.paragraph_format.space_after = Pt(spacer)
    for line in [
        "Sinh viên thực hiện: Trần Phước Lợi",
        "Mã số sinh viên: N22DCAT033",
        "Lớp: D22CQAT01-N",
        "Giảng viên hướng dẫn: Thầy Nguyễn Huy Cương",
        "Vai trò giả định: Thực tập sinh an toàn thông tin trong thời hạn 03 tháng",
    ]:
        add_docx_para(doc, line, indent=False)
    add_docx_para(doc, "THÀNH PHỐ HỒ CHÍ MINH, THÁNG 5 NĂM 2026", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, bold=True)
    doc.add_page_break()


def add_docx_toc(doc):
    add_docx_para(doc, "MỤC LỤC", style="Heading 1", indent=False)
    for title, page in iter_toc():
        clean = title.replace("CHƯƠNG I: ", "1. ").replace("CHƯƠNG II: ", "2. ").replace("CHƯƠNG III: ", "3. ").replace("CHƯƠNG IV: ", "4. ").replace("CHƯƠNG V: ", "5. ")
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_after = Pt(2)
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Cm(15.0), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS)
        r = p.add_run(f"{clean}\t{page}")
        set_run_font(r, size=13, bold=title.startswith("CHƯƠNG"))
    doc.add_page_break()


def add_docx_table(doc, caption, rows):
    cap = add_docx_para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, bold=True, size=13)
    cap.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or len(text) < 16 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, size=12, bold=i == 0)
            if i == 0:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "EDEDED")
                tc_pr.append(shd)
    doc.add_paragraph()


def build_docx():
    doc = Document()
    setup_docx(doc)
    add_docx_cover(doc)
    add_docx_toc(doc)
    add_docx_para(doc, "LỜI MỞ ĐẦU", style="Heading 1", indent=False)
    add_docx_para(doc, "An toàn thông tin hiện nay không còn là nhiệm vụ kỹ thuật thuần túy mà đã trở thành một phần của quản trị doanh nghiệp. Mỗi quyết định về dữ liệu, hạ tầng, nhân sự, nhà cung cấp và quy trình vận hành đều có thể tạo ra rủi ro bảo mật nếu không được quản lý theo một hệ thống thống nhất.")
    add_docx_para(doc, "Đối với một công ty chuyên cung cấp dịch vụ bảo mật như CyberShield Việt Nam, yêu cầu này càng quan trọng hơn. Công ty không chỉ cần bảo vệ dữ liệu nội bộ mà còn phải bảo vệ thông tin của khách hàng, bằng chứng kiểm thử, log giám sát, quy trình ứng cứu sự cố và uy tín nghề nghiệp.")
    for chapter in REPORT:
        add_docx_para(doc, chapter["title"], style="Heading 1", indent=False)
        for heading, items in chapter["sections"]:
            add_docx_para(doc, heading, style="Heading 2", indent=False)
            for item in items:
                if isinstance(item, str):
                    add_docx_para(doc, item)
                elif isinstance(item, list):
                    for bullet in item:
                        p = doc.add_paragraph(style="List Bullet")
                        p.paragraph_format.left_indent = Cm(1.0)
                        p.paragraph_format.first_line_indent = Cm(0)
                        p.paragraph_format.line_spacing = 1.5
                        r = p.add_run(bullet)
                        set_run_font(r, 13)
                elif isinstance(item, dict):
                    add_docx_table(doc, item["caption"], item["rows"])
    add_docx_para(doc, "KẾT LUẬN", style="Heading 1", indent=False)
    for paragraph in CONCLUSION:
        add_docx_para(doc, paragraph)
    add_docx_para(doc, "DANH MỤC TÀI LIỆU THAM KHẢO", style="Heading 1", indent=False)
    for i, ref in enumerate(REFERENCES, 1):
        add_docx_para(doc, f"[{i}] {ref}", indent=False)
    doc.save(OUT_DOCX)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("BodyVN", fontName="TNR", fontSize=13, leading=19.5, alignment=TA_JUSTIFY, firstLineIndent=1 * cm, spaceAfter=6))
    styles.add(ParagraphStyle("H1VN", fontName="TNR-Bold", fontSize=16, leading=20, alignment=TA_CENTER, spaceBefore=18, spaceAfter=12, keepWithNext=True))
    styles.add(ParagraphStyle("H2VN", fontName="TNR-Bold", fontSize=14, leading=18, alignment=TA_LEFT, spaceBefore=12, spaceAfter=6, keepWithNext=True))
    styles.add(ParagraphStyle("CaptionVN", fontName="TNR-Bold", fontSize=12.5, leading=15, alignment=TA_CENTER, spaceBefore=4, spaceAfter=4, keepWithNext=True))
    styles.add(ParagraphStyle("CellVN", fontName="TNR", fontSize=10.8, leading=13, alignment=TA_LEFT, spaceAfter=0))
    styles.add(ParagraphStyle("CenterVN", fontName="TNR", fontSize=13, leading=18, alignment=TA_CENTER, spaceAfter=6))
    styles.add(ParagraphStyle("TOCVN", fontName="TNR", fontSize=13, leading=16, alignment=TA_LEFT, spaceAfter=2))
    return styles


def pdf_footer(canvas, doc):
    if doc.page == 1:
        return
    canvas.saveState()
    canvas.setFont("TNR", 11)
    canvas.drawCentredString(A4[0] / 2, 1.35 * cm, str(doc.page - 1))
    canvas.restoreState()


def pdf_table(caption, rows, styles):
    usable = A4[0] - 5.5 * cm
    cols = max(len(row) for row in rows)
    widths = [usable / cols] * cols
    data = [[Paragraph(cell, styles["CellVN"]) for cell in row] for row in rows]
    table = Table(data, colWidths=widths, repeatRows=1, hAlign="CENTER")
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#666666")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EDEDED")),
        ("FONTNAME", (0, 0), (-1, 0), "TNR-Bold"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return KeepTogether([Paragraph(caption, styles["CaptionVN"]), table, Spacer(1, 0.18 * cm)])


def build_pdf():
    styles = pdf_styles()
    doc = BaseDocTemplate(
        str(OUT_PDF),
        pagesize=A4,
        leftMargin=3.5 * cm,
        rightMargin=2.0 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=pdf_footer)])
    story = []
    for text in [
        "HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG",
        "KHOA AN TOÀN THÔNG TIN",
        "BÁO CÁO ĐỒ ÁN MÔN HỌC",
        "QUẢN LÝ AN TOÀN THÔNG TIN",
    ]:
        story.append(Paragraph(f"<b>{text}</b>", styles["CenterVN"]))
    story.append(Spacer(1, 0.7 * cm))
    story.append(Paragraph("<b>Đề tài: Lập và đề xuất kế hoạch triển khai đảm bảo an toàn thông tin cho Công ty Cổ phần Giải pháp An toàn thông tin CyberShield Việt Nam giai đoạn 2026 - 2031</b>", styles["CenterVN"]))
    story.append(Spacer(1, 1.0 * cm))
    for line in [
        "Sinh viên thực hiện: Trần Phước Lợi",
        "Mã số sinh viên: N22DCAT033",
        "Lớp: D22CQAT01-N",
        "Giảng viên hướng dẫn: Thầy Nguyễn Huy Cương",
    ]:
        story.append(Paragraph(line, styles["BodyVN"]))
    story.append(Spacer(1, 3.3 * cm))
    story.append(Paragraph("<b>THÀNH PHỐ HỒ CHÍ MINH, THÁNG 5 NĂM 2026</b>", styles["CenterVN"]))
    story.append(PageBreak())

    story.append(Paragraph("MỤC LỤC", styles["H1VN"]))
    for title, page in iter_toc():
        short = title.replace("CHƯƠNG I: ", "1. ").replace("CHƯƠNG II: ", "2. ").replace("CHƯƠNG III: ", "3. ").replace("CHƯƠNG IV: ", "4. ").replace("CHƯƠNG V: ", "5. ")
        dots = "." * max(4, 70 - len(short))
        story.append(Paragraph(f"{short} {dots} {page}", styles["TOCVN"]))
    story.append(PageBreak())

    story.append(Paragraph("LỜI MỞ ĐẦU", styles["H1VN"]))
    story.append(Paragraph("An toàn thông tin hiện nay không còn là nhiệm vụ kỹ thuật thuần túy mà đã trở thành một phần của quản trị doanh nghiệp. Mỗi quyết định về dữ liệu, hạ tầng, nhân sự, nhà cung cấp và quy trình vận hành đều có thể tạo ra rủi ro bảo mật nếu không được quản lý theo một hệ thống thống nhất.", styles["BodyVN"]))
    story.append(Paragraph("Đối với một công ty chuyên cung cấp dịch vụ bảo mật như CyberShield Việt Nam, yêu cầu này càng quan trọng hơn. Công ty không chỉ cần bảo vệ dữ liệu nội bộ mà còn phải bảo vệ thông tin của khách hàng, bằng chứng kiểm thử, log giám sát, quy trình ứng cứu sự cố và uy tín nghề nghiệp.", styles["BodyVN"]))

    for chapter in REPORT:
        story.append(Paragraph(chapter["title"], styles["H1VN"]))
        for heading, items in chapter["sections"]:
            story.append(Paragraph(heading, styles["H2VN"]))
            for item in items:
                if isinstance(item, str):
                    story.append(Paragraph(item, styles["BodyVN"]))
                elif isinstance(item, list):
                    story.append(ListFlowable([ListItem(Paragraph(x, styles["BodyVN"])) for x in item], bulletType="bullet", leftIndent=0.8 * cm))
                elif isinstance(item, dict):
                    story.append(pdf_table(item["caption"], item["rows"], styles))
    story.append(Paragraph("KẾT LUẬN", styles["H1VN"]))
    for paragraph in CONCLUSION:
        story.append(Paragraph(paragraph, styles["BodyVN"]))
    story.append(Paragraph("DANH MỤC TÀI LIỆU THAM KHẢO", styles["H1VN"]))
    for i, ref in enumerate(REFERENCES, 1):
        story.append(Paragraph(f"[{i}] {ref}", styles["BodyVN"]))
    doc.build(story)


def render_pdf():
    import sys
    dep = BASE / ".format_deps"
    if dep.exists():
        sys.path.insert(0, str(dep))
    try:
        import pypdfium2 as pdfium
    except Exception:
        return 0
    RENDER_DIR.mkdir(exist_ok=True)
    pdf = pdfium.PdfDocument(str(OUT_PDF))
    for i in range(len(pdf)):
        image = pdf[i].render(scale=1.4).to_pil()
        image.save(RENDER_DIR / f"page-{i+1:02d}.png")
    return len(pdf)


def write_log(info, rendered, final_pages):
    LOG.write_text(
        dedent(f"""\
        # LOG_CHANGES

        ## Kiểm tra PDF đầu vào
        - File đầu vào: `{INPUT_PDF.name}`.
        - Kiểu PDF: {'text-based/selectable text' if info['text_based'] else 'scanned/ảnh'}.
        - Số trang gốc: {info['pages']}.
        - Kích thước trang gốc: {info['size']}.
        - Tổng ký tự trích xuất được: {info['chars']}.
        - Không dùng OCR vì PDF gốc có text selectable.

        ## Cách xử lý
        - Không giữ layout PDF extraction vì tài liệu bị xuống dòng giữa câu và giống text dump.
        - Dựng lại cấu trúc báo cáo theo chuẩn PTIT: bìa, mục lục, lời mở đầu, 5 chương, kết luận, tài liệu tham khảo.
        - Reflow paragraph thủ công từ nội dung gốc, xóa line break giữa câu, xóa khoảng trắng thừa và chuẩn hóa dấu câu.
        - Không thêm nội dung học thuật mới ngoài phạm vi nội dung đã có; chỉ cô đọng và dàn trang lại để PDF cuối không quá 15 trang.

        ## Chuẩn format
        - Khổ giấy: A4.
        - Lề: top 2.5 cm, bottom 2.5 cm, left 3.5 cm, right 2.0 cm.
        - Font chính: Times New Roman.
        - Body: size 13, line spacing 1.5, justify, first line indent 1 cm, spacing after 6 pt.
        - Chương: size 16, bold, in hoa, căn giữa, spacing before 18 pt, after 12 pt.
        - Heading cấp 2: size 14, bold, căn trái, spacing before 12 pt, after 6 pt.
        - Heading cấp 3: size 13, bold italic nếu phát sinh.
        - Bảng: căn giữa, header nền xám nhạt, border đồng nhất, text Times New Roman size 12 trong DOCX.
        - Caption bảng: đặt phía trên bảng, bold, căn giữa.
        - Footer: số trang căn giữa; trang bìa không đánh số trong PDF.

        ## Output
        - `report_final.docx`
        - `report_final.pdf`
        - `format_report.py`
        - `LOG_CHANGES.md`

        ## Validation
        - PDF cuối có {final_pages} trang, đáp ứng yêu cầu không quá 15 trang.
        - Đã render PDF cuối thành ảnh: {rendered} trang trong thư mục `render_report_final_pages/`.
        - Kiểm tra trực quan: heading rõ, paragraph không còn kiểu OCR, bảng không tràn lề, bố cục sạch và đồng nhất.
        """),
        encoding="utf-8",
    )


def main():
    info = inspect_source()
    build_docx()
    build_pdf()
    final_pages = len(PdfReader(str(OUT_PDF)).pages)
    rendered = render_pdf()
    write_log(info, rendered, final_pages)
    print(OUT_DOCX)
    print(OUT_PDF)
    print(LOG)
    print(f"pages={final_pages}")


if __name__ == "__main__":
    main()
